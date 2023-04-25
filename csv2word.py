import csv
import pandas as pd
import numpy as np
import docx
# define path+name of source file
in_path=r"C:\Users\Vernon\Desktop\critical_writing\硕士研究生英语-硕士研究生英语32班-随堂练习统计详情.xlsx"
# define main path of output writing 1
out1_path=r"C:/Users/Vernon/Desktop/critical_writing/class_32/writing_1/"
# define main path of output writing 2
out2_path=r"C:/Users/Vernon/Desktop/critical_writing/class_32/writing_2/"

def excel_one_line_to_list(path):
    '''
    :param path: path+name of source file(excel file)
    :return: name: list of student name
             result1: list of writing 1
             result2: list of writing 2
    '''
    df=pd.read_excel(path)
    # df = pd.read_excel(path, usecols=[1],names=None)  # 读取项目名称列,不要列名
    df_li = df.values.tolist()
    df_li=df_li[5:]
    name=[]
    result1 = []
    result2=[]
    for s_li in df_li:
        name.append(s_li[0])
        result1.append(s_li[9])
        result2.append(s_li[10])

    N=len(name)
    name=[name[index] for index in range(N) if result1[index]==result1[index]  ]
    result1=[item for item in result1 if item==item]
    result2=[item for item in result2 if item==item]
    return name,result1,result2


if __name__ == '__main__':
    name,result1,result2=excel_one_line_to_list(in_path)
    N = len(name)
    # transfer writing 1 to word(.docx)
    for index in range(N):
        file = docx.Document()
        file.add_paragraph(result1[index])
        save_path=out1_path+name[index]+'_1.docx'
        file.save(save_path)
    del result1
    # transfer writing 2 to word(.docx)
    for index in range(N):
        file = docx.Document()
        file.add_paragraph(result2[index])
        save_path = out2_path + name[index] + '_2.docx'
        file.save(save_path)
